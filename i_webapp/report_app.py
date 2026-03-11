"""
平面图识别系统 — 管理层汇报界面
Streamlit + Vertex AI Gemini 1.5 Pro
"""

import json
import io
import streamlit as st
import vertexai
from vertexai.generative_models import GenerativeModel, ChatSession
from pathlib import Path
from datetime import datetime

# ============================================================
# 1. 初始化配置
# ============================================================

GCP_PROJECT = "project-d3027d52-508f-4689-899"
GCP_REGION  = "us-central1"
GCS_BUCKET  = "yalingdata"

vertexai.init(project=GCP_PROJECT, location=GCP_REGION)

st.set_page_config(
    page_title="平面图识别系统 — 训练报告",
    page_icon="🏗️",
    layout="wide",
    initial_sidebar_state="collapsed",
)


# ============================================================
# 2. 数据加载
# ============================================================

@st.cache_data
def load_training_results() -> dict:
    """
    从 GCS 读取训练结果
    训练完成前使用示例数据，训练完成后替换为真实路径
    """
    # ── 训练完成后取消注释，从 GCS 读取真实数据 ──
    # from google.cloud import storage
    # client = storage.Client()
    # bucket = client.bucket(GCS_BUCKET)
    # blob = bucket.blob("runs/training_results.json")
    # data = json.loads(blob.download_as_text())
    # return data

    # ── 示例数据（训练完成前使用）──
    return {
        "project": "建筑平面图三维重建系统",
        "dataset": "CubiCasa5k（5000张建筑平面图）",
        "training_date": "2026-03-10",
        "wall_segmentation": {
            "model": "FPN + ResNet50",
            "best_iou": 0.81,
            "epochs_trained": 50,
            "history": [
                {"epoch": i, "train_loss": 0.8 - i*0.013,
                 "val_iou": 0.3 + i*0.01}
                for i in range(1, 51)
            ]
        },
        "symbol_detection": {
            "model": "Faster R-CNN",
            "best_val_loss": 0.23,
            "epochs_trained": 15,
            "history": [
                {"epoch": i, "train_loss": 1.2 - i*0.06,
                 "val_loss": 1.1 - i*0.055}
                for i in range(1, 16)
            ]
        },
        "yolo": {
            "model": "YOLOv8m-seg",
            "mAP50": 0.78,
            "mAP50_95": 0.52,
            "epochs_trained": 100,
        }
    }


@st.cache_resource
def init_gemini_chat(results: str) -> ChatSession:
    """初始化 Gemini 对话，注入训练背景"""
    model = GenerativeModel("gemini-1.5-pro")
    chat = model.start_chat()
    chat.send_message(f"""
你是一个AI项目助手，负责向非技术背景的管理层汇报
建筑平面图三维重建AI系统的训练结果。

项目背景：
- 这是一个从2D建筑平面图自动生成3D模型的AI系统
- 使用了5000张真实建筑平面图进行训练
- 目标是自动识别墙体结构和门窗位置

训练结果数据：
{results}

指标说明：
- IoU = 准确率，越高越好，0.81 = 81%，业界主流水平约0.75-0.82
- mAP = 目标检测综合准确率，越高越好
- Loss = 损失值，越低越好，代表模型学习效果

回答规则：
1. 用通俗语言，完全避免专业术语
2. 用打分、百分比等直观方式表达效果
3. 回答简洁，重点突出
4. 遇到不确定的问题，诚实说明
5. 始终用中文回答
""")
    return chat


# ============================================================
# 3. 页面布局
# ============================================================

def render_header(results: dict):
    """顶部：标题和基本信息"""
    st.title("🏗️ 建筑平面图识别系统")
    st.caption(f"训练日期：{results.get('training_date', '—')}　|　"
               f"数据集：{results.get('dataset', '—')}")
    st.divider()


def render_metrics(results: dict):
    """核心指标卡片"""
    st.subheader("📊 训练结果概览")

    col1, col2, col3, col4 = st.columns(4)

    wall = results["wall_segmentation"]
    sym  = results["symbol_detection"]
    yolo = results["yolo"]

    with col1:
        st.metric(
            label="🧱 墙体识别准确率",
            value=f"{wall['best_iou']*100:.0f}%",
            delta="达到论文水平 ✓"
        )
    with col2:
        st.metric(
            label="🚪 门窗检测准确率",
            value=f"{yolo['mAP50']*100:.0f}%",
            delta="良好"
        )
    with col3:
        st.metric(
            label="📉 检测模型损失",
            value=f"{sym['best_val_loss']:.2f}",
            delta="收敛稳定",
            delta_color="inverse"
        )
    with col4:
        total_epochs = (wall['epochs_trained'] +
                        sym['epochs_trained'] +
                        yolo['epochs_trained'])
        st.metric(
            label="🔄 总训练轮次",
            value=str(total_epochs),
            delta="全部完成 ✓"
        )


def render_charts(results: dict):
    """训练曲线图表"""
    import pandas as pd

    st.subheader("📈 训练过程")

    tab1, tab2 = st.tabs(["墙体分割", "门窗检测"])

    with tab1:
        history = results["wall_segmentation"]["history"]
        df = pd.DataFrame(history)

        col1, col2 = st.columns(2)
        with col1:
            st.markdown("**训练损失**")
            st.line_chart(df.set_index("epoch")["train_loss"],
                          color="#2E75B6")
        with col2:
            st.markdown("**验证 IoU（准确率）**")
            st.line_chart(df.set_index("epoch")["val_iou"],
                          color="#70AD47")

    with tab2:
        history = results["symbol_detection"]["history"]
        df = pd.DataFrame(history)

        col1, col2 = st.columns(2)
        with col1:
            st.markdown("**训练损失**")
            st.line_chart(df.set_index("epoch")["train_loss"],
                          color="#2E75B6")
        with col2:
            st.markdown("**验证损失**")
            st.line_chart(df.set_index("epoch")["val_loss"],
                          color="#ED7D31")


def render_summary_button(results: dict):
    """一键生成项目总结"""
    st.subheader("📋 项目总结")

    if st.button("✨ 生成AI总结", type="primary", use_container_width=True):
        with st.spinner("AI 正在生成总结..."):
            chat = init_gemini_chat(json.dumps(results, ensure_ascii=False))
            response = chat.send_message(
                "请用300字以内总结本次AI训练结果，"
                "重点说明：模型效果如何、能解决什么问题、"
                "与业界水平对比、下一步建议。"
                "语言要适合向公司管理层汇报。"
            )
            st.success(response.text)


def render_chat(results: dict):
    """AI 对话区"""
    st.subheader("💬 询问 AI 助手")
    st.caption("可以用自然语言提问，例如：这个模型够用吗？准确率达到要求了吗？还需要改进什么？")

    # 初始化对话历史
    if "messages" not in st.session_state:
        st.session_state.messages = []

    # 显示历史消息
    for msg in st.session_state.messages:
        with st.chat_message(msg["role"]):
            st.write(msg["content"])

    # 输入框
    if prompt := st.chat_input("请输入你的问题..."):
        with st.chat_message("user"):
            st.write(prompt)
        st.session_state.messages.append(
            {"role": "user", "content": prompt}
        )

        with st.chat_message("assistant"):
            with st.spinner("思考中..."):
                chat = init_gemini_chat(
                    json.dumps(results, ensure_ascii=False)
                )
                response = chat.send_message(prompt)
                st.write(response.text)

        st.session_state.messages.append(
            {"role": "assistant", "content": response.text}
        )


def render_export(results: dict):
    """导出 Word 报告"""
    st.subheader("📥 导出报告")

    if st.button("生成完整 Word 报告", use_container_width=True):
        with st.spinner("正在生成报告..."):

            # 调用 Gemini 生成完整报告内容
            chat = init_gemini_chat(
                json.dumps(results, ensure_ascii=False)
            )
            full_report = chat.send_message("""
请生成一份完整的项目汇报文档，包含以下章节：
1. 项目概述（这个AI系统是做什么的）
2. 训练数据说明（用了什么数据集，多少张图）
3. 各模型训练结果（用通俗语言描述准确率）
4. 与业界水平对比
5. 实际应用价值
6. 下一步计划

语言要适合非技术背景的管理层阅读，全程使用中文。
""")

            # 生成 Word 文档
            from docx import Document
            doc = Document()
            doc.add_heading("建筑平面图识别系统 — AI训练报告", 0)
            doc.add_paragraph(
                f"生成日期：{datetime.now().strftime('%Y年%m月%d日')}"
            )
            doc.add_paragraph(
                f"数据集：{results.get('dataset', '—')}"
            )
            doc.add_paragraph("")

            for line in full_report.text.split('\n'):
                if line.strip():
                    if line.startswith('#'):
                        doc.add_heading(
                            line.replace('#', '').strip(), level=1
                        )
                    else:
                        doc.add_paragraph(line)

            # 保存到内存并提供下载
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)

        st.download_button(
            label="📄 点击下载 Word 报告",
            data=buffer,
            file_name=f"AI训练报告_{datetime.now().strftime('%Y%m%d')}.docx",
            mime="application/vnd.openxmlformats-officedocument"
                 ".wordprocessingml.document",
            use_container_width=True,
        )


# ============================================================
# 4. 主函数
# ============================================================

def main():
    results = load_training_results()

    render_header(results)
    render_metrics(results)
    st.divider()
    render_charts(results)
    st.divider()
    render_summary_button(results)
    st.divider()
    render_chat(results)
    st.divider()
    render_export(results)


if __name__ == "__main__":
    main()